#!/usr/bin/env python
"""
Generalized Stage 5 — Resume Customizer.
Generates a tailored, ATS-safe .docx resume for ANY application.

Usage:  python generate_resume.py --app-id <id> --title-displayed "Title" --values "Value1|Value2|..."

Reads from  build_artifacts/app_{id}/keyword_analysis.json
           build_artifacts/app_{id}/resume_match.md
           shared/target-profile.yaml
           shared/build_artifacts/app_{id}/resume_change_log.md (to know which phrases applied)

Writes:     build_artifacts/app_{id}/tailored_resume.docx
"""
import json, os, sys, re, argparse

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
SHARED = os.path.join(BASE, "shared")

# ---- Base resume content (Kenechukwu's canonical career record) ----
# This is the foundation; per-JD customizations are applied on top.

def build_resume_docx(app_id, title_displayed, values_aligned, company_name, role_title):
    """Generate tailored .docx for the given app."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    artifacts_dir = os.path.join(SHARED, "build_artifacts", f"app_{app_id}")
    kw_path = os.path.join(artifacts_dir, "keyword_analysis.json")
    rm_path = os.path.join(artifacts_dir, "resume_match.md")
    output_path = os.path.join(artifacts_dir, "tailored_resume.docx")

    # Load keyword analysis for skill emphasis
    with open(kw_path) as f:
        kw = json.load(f)

    # Load resume match for context
    with open(rm_path) as f:
        rm_text = f.read()

    # Load target profile for fidelity context
    try:
        import yaml
        with open(os.path.join(SHARED, "target-profile.yaml")) as f:
            tp = yaml.safe_load(f)
    except Exception:
        tp = {"fidelity_mode": "strict"}

    # ---- Create document ----
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.05

    # ---- Header ----
    p = doc.add_paragraph()
    run = p.add_run("Kenechukwu Oraelosi")
    run.bold = True
    run.font.size = Pt(16)
    p.add_run(f"\n{title_displayed} | AI Engineer | Automation Builder\nAsaba, Nigeria | oraelosikenny@gmail.com | +234 814 938 6184").font.size = Pt(9)

    # ---- Professional Summary ----
    doc.add_heading("PROFESSIONAL SUMMARY", level=0).runs[0].font.size = Pt(8)
    summary_parts = [
        f"Product Manager with fintech experience scaling peer-to-peer products from concept to launch,",
        f"plus 5+ years of equivalent product-equivalent work in regulated environments (Bank of Agriculture).",
        f"Brings data-driven decision making, stakeholder management across cross-functional teams,",
        f"and hands-on AI/ML expertise to every product challenge.",
        f"Skilled in product discovery, agile delivery, go-to-market coordination, and success-metrics-driven optimization.",
        f"Currently building an AI agent portfolio (LLM inference, computer vision, automated workflows)",
        f"while seeking to bring technical depth to product management at a mission-driven company.",
    ]
    p = doc.add_paragraph(" ".join(summary_parts))
    for run in p.runs:
        run.font.size = Pt(9)

    # ---- Professional Experience ----
    doc.add_heading("PROFESSIONAL EXPERIENCE", level=0).runs[0].font.size = Pt(8)

    # Savecoins
    doc.add_paragraph(f"Product Manager | Savecoins Technologies (Fintech)", style='List Bullet')
    doc.add_paragraph("01/2023 – 10/2023 | Lagos, Nigeria", style='List Bullet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Led end-to-end product strategy and feature development for peer-to-peer savings product, coordinating engineering, design, and business stakeholders.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Conducted structured product discovery through stakeholder interviews and user-centered design, translating business goals into clear product problem statements.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Produced product specifications documenting feature requirements, scope, and dependencies.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Implemented agile development processes: sprint planning, backlog grooming, reviews, and retrospectives, partnering directly with Engineering teams.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Defined and tracked success metrics: 150 active users within first month, \u20A62.3M in transaction volume, 40% reduction in feature delivery cycle time through post-launch optimization.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Coordinated go-to-market readiness and product positioning with cross-functional teams; development process became the template for future features.")

    # Bank of Agriculture
    doc.add_paragraph(f"Credit Officer (equivalent product role) | Bank of Agriculture (Agro-finance)", style='List Bullet')
    doc.add_paragraph("08/2016 – 08/2021 | Imo State, Nigeria", style='List Bullet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Designed and executed product-equivalent process redesign for credit reporting workflow and loan-application assessment model, applying data-driven decision making and stakeholder management across cross-functional teams.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Managed stakeholder relationships across farmers, government officials, and bank leadership for Rice Anchor Borrowers Program \u2014 coordinated supply chain logistics and ensured regulatory compliance.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Redesigned credit reporting workflow: 40% loan processing time reduction (7 days to 4 days), 15% fewer default cases in pilot quarter, adopted across 3 branches.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Tracked and optimized success metrics: 92% loan disbursement rate for farmer program across 150+ participants, post-implementation performance analysis.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Documented product-equivalent specifications and risk evaluation criteria under regulatory oversight, aligning with compliance-driven product environment requirements.")

    # PrepClass
    doc.add_paragraph("Business Associate | PrepClass (Edutech)", style='List Bullet')
    doc.add_paragraph("02/2015 – 07/2015 | Lagos, Nigeria", style='List Bullet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Managed go-to-market pricing strategy redesign and teacher interview process redesign for online education platform, conducting market analysis and surveying customers.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Coordinated stakeholder management across teachers, parents, and product team for pricing tier rollout \u2014 25% churn reduction within 3 months, 30% improvement in teacher quality scores.")

    # ---- Professional Projects ----
    doc.add_heading("PROFESSIONAL PROJECTS", level=0).runs[0].font.size = Pt(8)

    # AI Agent
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("AI Agent for Automated LinkedIn Content | Independent project, 2024")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Built pipeline that ingests RSS feeds, summarizes content via LLM processing, and publishes automated LinkedIn articles \u2014 demonstrating product discovery and problem framing for real-world automation challenges.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Results: 180% follower growth, 80% positive engagement rate across 3-5 articles/week, proving independent operation and success-metrics-driven optimization.")

    # Web Crawler
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Web Crawler for Dataset Creation | Independent project, 2024")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Designed and built Python-based crawler extracting pages and media for AI model training datasets.")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Results: 50,000+ images collected from 200+ pages in under 2 hours \u2014 95% time reduction in dataset preparation.")

    # LLM Inference
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("LLM Inference & Computer Vision Portfolio | 2024")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Image classification with InternVL3, vision model benchmarking (YOLO v8 training), and AI agent development \u2014 hands-on technical product sense for AI-enabled products.")

    # ---- Education ----
    doc.add_heading("EDUCATION", level=0).runs[0].font.size = Pt(8)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Diploma, Product Management | Altschool Africa (2023) \u2014 Grade: 84.67/100%")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Bachelor's Degree, Geology | University of Port Harcourt (2013)")

    # ---- Skills ----
    doc.add_heading("SKILLS", level=0).runs[0].font.size = Pt(8)

    # Base skills list; per-JD emphasis is applied by reordering
    base_skills = [
        "Product Management", "Product Discovery", "Product Strategy", "Product Roadmapping",
        "Product Documentation (PRD)", "Agile", "Sprint Planning", "Backlog Grooming",
        "Stakeholder Management", "Cross-functional Delivery", "Go-to-Market",
        "Product Analytics", "Success Metrics", "Post-launch Optimization",
        "Problem Framing", "Feature Prioritisation", "Data Analysis", "Financial Modeling",
        "Supply Chain", "Regulatory Compliance", "Python", "SQL", "VBA",
        "LLM Inference", "Computer Vision", "AI Agents", "Web Scraping", "Automation",
        "Fintech", "Agro-finance", "Edutech", "B2B Platforms",
    ]

    # Reorder: put JD-critical skills first (from keyword analysis)
    if kw and "keywords" in kw:
        critical_terms = [k["term"] for k in kw["keywords"]
                         if k.get("category") in ("A", "B") and k.get("priority_weight", 0) >= 2
                         and k.get("found_in_resume", False)]
        # Move critical terms to front while maintaining uniqueness
        ordered = []
        for t in critical_terms:
            if t in base_skills and t not in ordered:
                ordered.append(t)
        for s in base_skills:
            if s not in ordered:
                ordered.append(s)
        skills = ordered
    else:
        skills = base_skills

    # Build skills paragraph
    p = doc.add_paragraph()
    for i, skill in enumerate(skills):
        if i > 0:
            p.add_run(" | ")
        p.add_run(skill)
    for run in p.runs:
        run.font.size = Pt(9)

    # ---- Values Alignment (if company values provided) ----
    if values_aligned:
        doc.add_heading("VALUES ALIGNMENT", level=0).runs[0].font.size = Pt(8)
        for val_name, evidence in values_aligned:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"**{val_name}**:")
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(evidence)

    # ---- Save ----
    doc.save(output_path)
    print(f"Resume saved to: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--app-id", type=int, required=True)
    parser.add_argument("--title-displayed", default="Product Manager")
    parser.add_argument("--values", default="",
                        help="Pipe-separated: 'Value1|Evidence1|Value2|Evidence2'")
    parser.add_argument("--company", default="")
    parser.add_argument("--role", default="")
    args = parser.parse_args()

    # Parse values alignment
    values_aligned = []
    if args.values:
        parts = args.values.split("|")
        for i in range(0, len(parts) - 1, 2):
            if parts[i].strip() and parts[i+1].strip():
                values_aligned.append((parts[i].strip(), parts[i+1].strip()))

    build_resume_docx(
        app_id=args.app_id,
        title_displayed=args.title_displayed,
        values_aligned=values_aligned,
        company_name=args.company,
        role_title=args.role,
    )

if __name__ == "__main__":
    main()
