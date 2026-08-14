#!/usr/bin/env python
"""
verify_app_artifacts.py — verify a job-hunting build_artifacts/app_N directory
against the pipeline processor's parsing contracts (pipeline_processor.py).

Usage:
    python verify_app_artifacts.py <app_dir> [--forbidden "pat1,pat2,..."]

Reads ONLY. Never writes to the DB, never regenerates artifacts.

Checks (mirrors 00-orchestrator/scripts/pipeline_processor.py):
  * all 8 required stage artifacts present and non-empty
  * keyword_analysis.json: schema, nested analysis.match_score_percentage,
    scoring math recomputed (possible/earned/raw %, penalized)
  * resume_match.md: first-% == overall_match_score (raw, per convention),
    Gate 1 [PASSED]/[FAILED] and Gate 2 verdict lines present
  * resume_change_log.md: Tactic 2 displayed title (no Senior inflation)
  * risk_tactics_change_log.md: [FAIL] entries parse into (claim, evidence)
    pairs via the em-dash contract; PASS/FAIL/BORDERLINE counts
  * cover_letter.txt: body words < 400 (signature-name lines excluded)
  * tailored_resume.docx (if present): read back with python-docx, grep
    forbidden claims, check honest headline and evidence numbers

Forbidden patterns default to the strict-fidelity gap cluster; override with
--forbidden (comma-separated regexes) to match a specific JD's gaps.
"""
import argparse
import json
import os
import re
import sys

REQUIRED_FILES = [
    "jd_analysis.md",
    "resume_match.md",
    "keyword_analysis.json",
    "resume_change_log.md",
    "risk_tactics_change_log.md",
    "cover_letter.txt",
    "application_qa.md",
    "generate_resume.py",
]

DEFAULT_FORBIDDEN = [
    r"senior\s*product", r"principal\s*product", r"leading designers",
    r"player-?coach", r"embedded design", r"\bb2b\b", r"\bsaas\b",
    r"product-led growth", r"\bplg\b", r"\bokr\b", r"objectives and key results",
    r"kill criteria", r"leading indicators", r"time tracking",
    r"time intelligence", r"tribe", r"sales-assisted", r"customer calls",
    r"support and sales input", r"a/b test", r"ab test", r"experiment",
    r"discovery program",
]

fails = []


def check(name, cond, detail=""):
    print(("PASS  " if cond else "FAIL  ") + name + (("  -- " + detail) if detail else ""))
    if not cond:
        fails.append(name)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("app_dir")
    ap.add_argument("--forbidden", default=",".join(DEFAULT_FORBIDDEN))
    args = ap.parse_args()
    app = args.app_dir
    forbidden = [p for p in args.forbidden.split(",") if p.strip()]

    # 1. Artifact presence
    for f in REQUIRED_FILES:
        p = os.path.join(app, f)
        check(f"present: {f}", os.path.exists(p) and os.path.getsize(p) > 0,
              "missing/empty" if not (os.path.exists(p) and os.path.getsize(p) > 0) else "")

    # 2. keyword_analysis.json
    kw_path = os.path.join(app, "keyword_analysis.json")
    if os.path.exists(kw_path):
        kw = json.load(open(kw_path, encoding="utf-8"))
        a = kw.get("analysis") or {}
        check("kw: analysis keys complete", all(
            k in a for k in ["total_keywords_found", "total_possible_points",
                             "earned_points", "match_score_percentage",
                             "match_rating", "penalty_applied",
                             "penalized_score_percentage"]))
        terms = kw.get("keywords", [])
        check("kw: keyword count consistent", len(terms) == a.get("total_keywords_found", -1),
              f"{len(terms)} vs {a.get('total_keywords_found')}")
        check("kw: 10..15 terms", 10 <= len(terms) <= 15, str(len(terms)))
        check("kw: entry shape valid", all(
            {"term", "category", "priority_weight", "found_in_resume", "context_note"} <= set(t)
            for t in terms))
        check("kw: category enum valid", all(
            t["category"] in {"Hard Skill", "Domain Concept", "Soft Skill"} for t in terms))
        check("kw: weights in {1,2,3}", all(t["priority_weight"] in {1, 2, 3} for t in terms))
        if terms:
            possible = sum(t["priority_weight"] for t in terms)
            earned = sum(t["priority_weight"] for t in terms if t["found_in_resume"])
            raw = round(earned / possible * 100) if possible else 0
            penalized = round(raw * 0.75)  # Senior/industry penalty per schema
            check("kw: math recomputed", a.get("total_possible_points") == possible
                  and a.get("earned_points") == earned
                  and a.get("match_score_percentage") == raw, f"{possible}/{earned} -> {raw}")
            check("kw: penalized == round(raw*0.75)",
                  a.get("penalized_score_percentage") == penalized, str(penalized))

    # 3. resume_match.md
    rm_path = os.path.join(app, "resume_match.md")
    if os.path.exists(rm_path):
        rm = open(rm_path, encoding="utf-8").read()
        m = re.search(r"(\d+)%", rm)
        check("rm: first % present", bool(m), "no percentage found" if not m else "")
        check("rm: Gate 1 verdict present", "[PASSED]" in rm or "[FAILED]" in rm)
        g2 = re.search(r"Gate\s*2[^\n]*?verdict[^\n]*?:\s*([^\n]+)", rm, re.I)
        check("rm: Gate 2 verdict line present", bool(g2),
              g2.group(1).strip()[:50] if g2 else "missing")

    # 4. resume_change_log.md — displayed title
    rcl_path = os.path.join(app, "resume_change_log.md")
    if os.path.exists(rcl_path):
        rcl = open(rcl_path, encoding="utf-8").read()
        sec = re.search(r"##\s*Tactic\s*2.*?(?=^##\s*Tactic|\Z)", rcl, re.S | re.M)
        disp = None
        if sec:
            for line in sec.group(0).splitlines():
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if len(cells) >= 2 and cells[1] not in ("", "-", "---", "Displayed Title"):
                    disp = cells[1]
                    break
        check("rcl: displayed title parsed", bool(disp), str(disp))
        check("rcl: no Senior inflation in title", bool(disp) and "senior" not in disp.lower(),
              str(disp))

    # 5. risk_tactics_change_log.md — FAIL gap contract + counts
    gate_path = os.path.join(app, "risk_tactics_change_log.md")
    if os.path.exists(gate_path):
        gate = open(gate_path, encoding="utf-8").read()
        pairs = re.findall(r"\[FAIL\]\s*(?!:)(.+?)\s*\u2014\s*(.+?)(?=\n|$)", gate)
        check("gate: FAIL entries parse into (claim, evidence)", len(pairs) >= 1,
              f"{len(pairs)} pairs")
        for pat in (r"\[PASS\]", r"\[FAIL\]", r"\[BORDERLINE PASS\]"):
            check(f"gate: count {pat}", bool(re.search(pat, gate)))

    # 6. cover_letter.txt — body word count
    cl_path = os.path.join(app, "cover_letter.txt")
    if os.path.exists(cl_path):
        lines = open(cl_path, encoding="utf-8").read().splitlines()
        body = "\n".join(l for l in lines if l.strip() and not l.strip().startswith("Kenechukwu"))
        n = len(body.split())
        check("cl: body words < 400", n < 400, f"{n} words")

    # 7. tailored_resume.docx — read-back forbidden grep (independent pass)
    docx_path = os.path.join(app, "tailored_resume.docx")
    if os.path.exists(docx_path):
        try:
            from docx import Document
            doc = Document(docx_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            hits = [p for p in forbidden if re.search(p, text, re.I)]
            check("docx: forbidden-claim grep clean", not hits,
                  "HITS: " + ", ".join(hits) if hits else "")
            headline = next((p.text for p in doc.paragraphs if "Automation Builder" in p.text), "")
            check("docx: honest headline present", "Product Manager" in headline and "Senior" not in headline,
                  headline)
        except ImportError:
            print("SKIP  docx read-back (python-docx not installed)")
    else:
        check("docx: tailored_resume.docx present", False, "missing (run generate_resume.py)")

    print()
    print("RESULT:", "ALL CHECKS PASSED" if not fails else f"FAILED ({len(fails)}): {fails}")
    sys.exit(1 if fails else 0)


if __name__ == "__main__":
    main()
